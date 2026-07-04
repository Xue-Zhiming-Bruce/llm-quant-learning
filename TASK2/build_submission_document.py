"""Build the formatted TASK2 submission document.

The script creates a DOCX source file. Use the documents renderer to export it
to PDF and inspect page images before submission.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
FIGURE_DIR = BASE_DIR / "figures"
DOCX_PATH = BASE_DIR / "薛智鸣TASK2.docx"

FONT_NAME = "Songti SC"
BODY_SIZE = Pt(10.5)


def set_run_font(run, size: Pt = BODY_SIZE, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), FONT_NAME)


def format_paragraph(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line: bool = False,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    if first_line:
        fmt.first_line_indent = Pt(21)
    else:
        fmt.first_line_indent = Pt(0)
    fmt.keep_with_next = keep_with_next
    for run in paragraph.runs:
        set_run_font(run)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(width_in * 1440)))
    width.set(qn("w:type"), "dxa")
    tc_pr.append(width)


def set_table_width(table, width_in: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_font(run, BODY_SIZE, bold=True)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, BODY_SIZE, bold=True)
    format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_with_next=True)


def add_body(doc: Document, text: str, *, first_line: bool = True) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)
    format_paragraph(paragraph, first_line=first_line)


def add_formula(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run)
        format_paragraph(paragraph, first_line=False)
        paragraph.paragraph_format.left_indent = Pt(21)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=True)
    format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, sum(widths))

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_width(header_cells[i], widths[i])
        set_cell_shading(header_cells[i], "EDEDED")
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = header_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(str(header))
        set_run_font(run, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run("" if pd.isna(value) else str(value))
            set_run_font(run)

    spacer = doc.add_paragraph()
    format_paragraph(spacer, first_line=False)


def pct_position(row: pd.Series) -> str:
    value = row["bb_percent_b"]
    if pd.isna(value):
        return ""
    if value >= 1:
        return "上轨上方"
    if value <= 0:
        return "下轨下方"
    if value >= 0.5:
        return "中上区间"
    return "中下区间"


def fmt_num(value: object, digits: int = 2) -> str:
    if pd.isna(value):
        return ""
    return f"{float(value):.{digits}f}"


def build_document() -> Path:
    required_paths = [
        OUTPUT_DIR / "diagnostics_summary.csv",
        OUTPUT_DIR / "descriptive_statistics.csv",
        OUTPUT_DIR / "000001.SZ_indicators.csv",
        OUTPUT_DIR / "600031.SH_indicators.csv",
        FIGURE_DIR / "000001.SZ_technical_indicators.png",
        FIGURE_DIR / "600031.SH_technical_indicators.png",
    ]
    missing = [str(path) for path in required_paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing analysis outputs: " + ", ".join(missing))

    diagnostics = pd.read_csv(OUTPUT_DIR / "diagnostics_summary.csv")
    descriptive = pd.read_csv(OUTPUT_DIR / "descriptive_statistics.csv")
    indicator_frames = [
        pd.read_csv(OUTPUT_DIR / "000001.SZ_indicators.csv"),
        pd.read_csv(OUTPUT_DIR / "600031.SH_indicators.csv"),
    ]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_title(doc, "TASK2：基础诊断与技术指标分析")
    add_body(doc, "姓名：薛智鸣", first_line=False)
    add_body(doc, "本文基于 TASK2 目录下已存储的股价日线数据，完成数据诊断、RSI、MACD、布林带和扩展指标 ATR 的计算，并给出可视化结果及解读。", first_line=True)

    add_heading(doc, "一、数据基础诊断分析")
    add_body(doc, "本次读取两份行情数据文件，字段包括开盘价、最高价、最低价、收盘价、昨收价、涨跌额、涨跌幅、成交量和成交额等。计算指标前先将交易日期转换为日期类型，并按交易日升序排列。", first_line=True)

    diag_rows = []
    for _, row in diagnostics.iterrows():
        diag_rows.append(
            [
                row["stock_name"],
                row["ts_code"],
                int(row["rows"]),
                f"{row['start_date']}至{row['end_date']}",
                int(row["total_missing_values"]),
                int(row["duplicate_trade_dates"]),
            ]
        )
    add_caption(doc, "表1 数据完整性诊断结果")
    add_table(
        doc,
        ["股票", "代码", "样本数", "日期范围", "缺失值", "重复交易日"],
        diag_rows,
        [1.0, 1.0, 0.75, 2.0, 0.8, 0.95],
    )
    add_body(doc, "诊断结果显示，两份数据均不存在缺失值和重复交易日，能够直接用于后续描述性统计和技术指标计算。需要说明的是，平安集团文件中的代码字段为 000001.SZ，本文在表格和图形中保留源文件名称与真实代码字段。", first_line=True)

    desc_rows = []
    for stock_name in diagnostics["stock_name"]:
        close_row = descriptive[(descriptive["stock_name"] == stock_name) & (descriptive["column"] == "close")].iloc[0]
        pct_row = descriptive[(descriptive["stock_name"] == stock_name) & (descriptive["column"] == "pct_chg")].iloc[0]
        desc_rows.append(
            [
                stock_name,
                fmt_num(close_row["mean"], 2),
                fmt_num(close_row["std"], 2),
                fmt_num(close_row["min"], 2),
                fmt_num(close_row["max"], 2),
                fmt_num(pct_row["mean"], 2),
                fmt_num(pct_row["std"], 2),
            ]
        )
    add_caption(doc, "表2 主要描述性统计量")
    add_table(
        doc,
        ["股票", "收盘均值", "收盘标准差", "收盘最小值", "收盘最大值", "涨跌幅均值(%)", "涨跌幅标准差(%)"],
        desc_rows,
        [0.9, 0.85, 0.95, 0.95, 0.95, 1.05, 1.05],
    )
    add_body(doc, "从描述性统计看，三一重工样本期内收盘价均值为 18.25，收盘价标准差为 1.36；平安集团文件样本期内收盘价均值为 10.95，收盘价标准差为 0.84。涨跌幅标准差反映了日度波动水平，三一重工约为 1.68%，平安集团文件约为 1.49%。", first_line=True)

    add_heading(doc, "二、RSI、MACD 和布林带的计算方法及作用")
    add_body(doc, "RSI 是相对强弱指标，用于衡量一段时间内上涨力度与下跌力度的相对关系，常用周期为 14。计算时先得到每日上涨幅度 Gain 和下跌幅度 Loss，再用 Wilder 平滑方法计算平均上涨和平均下跌。", first_line=True)
    add_formula(
        doc,
        [
            "Gain_t=max(Close_t-Close_(t-1),0)，Loss_t=max(Close_(t-1)-Close_t,0)",
            "RS=AvgGain/AvgLoss",
            "RSI=100-100/(1+RS)",
        ],
    )
    add_body(doc, "RSI 取值范围为 0 到 100，通常 RSI 高于 70 表示短期偏热，低于 30 表示短期偏冷。它适合辅助观察超买超卖、动能变化和背离，但在强趋势中可能长时间停留在高位或低位。", first_line=True)

    add_body(doc, "MACD 是指数平滑异同移动平均线，常用参数为 12、26、9。它先计算 12 日 EMA 和 26 日 EMA 的差值，再对差值计算 9 日 EMA 作为信号线。", first_line=True)
    add_formula(
        doc,
        [
            "EMA_t=alpha*Close_t+(1-alpha)*EMA_(t-1)，alpha=2/(N+1)",
            "MACD Line=EMA_12-EMA_26",
            "Signal Line=EMA_9(MACD Line)，Histogram=MACD Line-Signal Line",
        ],
    )
    add_body(doc, "MACD 主要用于判断趋势方向和动能变化。MACD 线上穿信号线通常被称为金叉，下穿信号线称为死叉；柱状图扩大说明短期动能增强，收缩说明动能减弱。", first_line=True)

    add_body(doc, "布林带由中轨、上轨和下轨组成，常用参数为 20 日均线和 2 倍标准差。它通过均线和波动率共同描述价格的相对高低。", first_line=True)
    add_formula(
        doc,
        [
            "Middle Band=SMA_20",
            "Upper Band=SMA_20+2*Std_20",
            "Lower Band=SMA_20-2*Std_20",
            "Bandwidth=(Upper Band-Lower Band)/Middle Band，%B=(Close-Lower Band)/(Upper Band-Lower Band)",
        ],
    )
    add_body(doc, "布林带常用于观察价格相对位置、波动率收缩或扩张，以及突破和均值回归机会。价格触及上下轨不代表一定反转，需要结合趋势、成交量或其他指标确认。", first_line=True)

    add_heading(doc, "三、Python 实现与指标结果")
    add_body(doc, "Python 脚本 analyze_technical_indicators.py 完成了数据加载、缺失值检查、描述性统计、RSI、MACD、布林带和 ATR 的计算，并将指标明细保存为 CSV 文件，将综合可视化图形保存到 figures 目录。", first_line=True)

    latest_rows = []
    latest_by_code = {}
    for frame in indicator_frames:
        frame["trade_date"] = pd.to_datetime(frame["trade_date"])
        row = frame.iloc[-1].copy()
        latest_by_code[row["ts_code"]] = row
        latest_rows.append(
            [
                row["stock_name"],
                row["trade_date"].date().isoformat(),
                fmt_num(row["close"], 2),
                fmt_num(row["rsi_14"], 2),
                fmt_num(row["macd_hist"], 4),
                pct_position(row),
                fmt_num(row["atr_14"], 4),
            ]
        )
    add_caption(doc, "表3 最新交易日技术指标摘要")
    add_table(
        doc,
        ["股票", "日期", "收盘价", "RSI14", "MACD柱", "布林位置", "ATR14"],
        latest_rows,
        [0.9, 1.15, 0.8, 0.75, 0.85, 1.0, 0.8],
    )
    add_body(doc, "从最新交易日指标看，两只股票的 RSI 均在 50 附近偏上，说明短期动能并未进入明显超买或超卖区域。平安集团文件对应代码 000001.SZ 的 MACD 柱略为负，显示短期动能较前期有所减弱；三一重工的 MACD 柱为正，显示短期修复动能仍在延续。", first_line=True)

    doc.add_page_break()
    add_caption(doc, "图1 平安集团文件（000001.SZ）RSI、MACD、布林带和 ATR 综合图")
    doc.add_picture(str(FIGURE_DIR / "000001.SZ_technical_indicators.png"), width=Inches(6.15))
    add_body(doc, "图1显示，000001.SZ 在 2024 年 10 月附近出现快速上行，布林带明显扩张，ATR 同步抬升，说明价格波动显著放大。2025 年 4 月附近 RSI 接近或跌破 30 后价格逐步回稳；截至最后交易日，RSI 为 56.39，MACD 柱为 -0.0185，价格位于布林带中上区间，说明整体尚未进入明显超买，但短期动能略有回落。", first_line=True)

    doc.add_page_break()
    add_caption(doc, "图2 三一重工（600031.SH）RSI、MACD、布林带和 ATR 综合图")
    doc.add_picture(str(FIGURE_DIR / "600031.SH_technical_indicators.png"), width=Inches(6.15))
    add_body(doc, "图2显示，三一重工在 2025 年 3 月前后价格快速上行，RSI 一度高于 70，MACD 位于高位且布林带扩张，反映出较强上涨动能和较高波动。随后 4 月至 6 月价格进入调整，ATR 逐步下降，说明波动收敛。截至最后交易日，RSI 为 55.65，MACD 柱为 0.0951，价格处于布林带中上区间，短期修复迹象较明显。", first_line=True)

    add_heading(doc, "四、扩展指标：ATR")
    add_body(doc, "除 RSI、MACD 和布林带外，量化分析中还常见均线 MA/EMA、动量 Momentum/ROC、KDJ、ADX、OBV、VWAP、波动率 Volatility 和 MFI 等指标。本文选取 ATR 作为扩展指标。", first_line=True)
    add_body(doc, "ATR 即平均真实波幅，用于衡量市场真实波动幅度，不直接判断涨跌方向。真实波幅 TR 的计算同时考虑当日最高价、最低价以及前一日收盘价，能反映跳空造成的风险。", first_line=True)
    add_formula(
        doc,
        [
            "TR_t=max(High_t-Low_t, abs(High_t-Close_(t-1)), abs(Low_t-Close_(t-1)))",
            "ATR_t=(ATR_(t-1)*13+TR_t)/14",
        ],
    )
    add_body(doc, "ATR 的作用主要包括动态止损、仓位控制和波动率过滤。例如可以用买入价减去 2 倍 ATR 作为止损参考，使止损距离随市场波动自动调整，避免固定止损在高波动阶段过窄、在低波动阶段过宽。", first_line=True)

    add_heading(doc, "五、结论")
    add_body(doc, "本次作业完成了两只股票日线数据的基础诊断、描述性统计、三类核心技术指标计算与可视化，并扩展计算了 ATR。总体来看，指标计算结果能够从动能、趋势和波动三个角度补充观察价格走势；但技术指标只能作为量化策略中的特征或信号来源，实际使用时仍需结合交易成本、回测检验和风险控制。", first_line=True)

    add_heading(doc, "参考资料")
    references = [
        "Investopedia, RSI: https://www.investopedia.com/terms/r/rsi.asp",
        "Investopedia, MACD: https://www.investopedia.com/terms/m/macd.asp",
        "Wikipedia, Bollinger Bands: https://en.wikipedia.org/wiki/Bollinger_Bands",
        "Wikipedia, Average True Range: https://en.wikipedia.org/wiki/Average_true_range",
    ]
    for item in references:
        add_body(doc, item, first_line=False)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        footer.paragraph_format.line_spacing = 1.5
        run = footer.add_run("薛智鸣TASK2")
        set_run_font(run)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
