"""Build the formally formatted TASK6 submission DOCX/PDF from analysis_report.md."""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "analysis_report.md"
DOCX_OUT = BASE_DIR / "薛智鸣TASK6.docx"
PDF_OUT = BASE_DIR / "薛智鸣TASK6.pdf"
# macOS's installed simplified-Chinese Song typeface. It renders as 宋体-简
# and avoids LibreOffice's invalid SimSun -> Times New Roman substitution.
FONT = "Songti SC"
BODY_SIZE = 10.5
BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)


TABLE_TITLES = [
    "常见机器学习量化自变量因子",
    "常见机器学习量化应变量",
    "测试集各季度策略收益率",
    "模型测试集预测效果",
    "策略回测核心指标",
]

FIGURE_INTERPRETATIONS = {
    "测试集累计净值": (
        "图1显示，测试期内三种机器学习策略的净值整体均高于全市场等权基准。"
        "随机森林在2022年第一季度反弹时取得较高收益，期末净值约为1.005；尽管测试期很短，"
        "其下行幅度和最终表现仍优于岭回归与单棵决策树。"
    ),
    "各季度收益率": (
        "图2显示，三种模型在2021年第四季度和2022年第二季度均出现绝对亏损，但亏损幅度小于等权基准；"
        "在2022年第一季度，模型策略均取得正收益。由此可见，测试期超额收益主要来自下跌季度的相对抗跌能力。"
    ),
    "模型预测效果对比": (
        "图3显示，随机森林的Rank IC最高，说明其对股票下一季度收益的截面排序能力最强；"
        "三种模型的方向准确率均较低，表明预测收益方向仍然困难。策略构建因此更适合使用相对排名，"
        "不宜把预测值直接解释为精确收益。"
    ),
}


def set_run_font(run, size=BODY_SIZE, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_paragraph(paragraph, *, justified=True, first_indent=True, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    fmt.keep_with_next = keep_next
    if first_indent:
        fmt.first_line_indent = Pt(BODY_SIZE * 2)
    if justified:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=9, color=GRAY)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clean_inline(text):
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\log", "log")
    return text


def add_rich_text(paragraph, text, size=BODY_SIZE):
    # Preserve bold emphasis from Markdown while using a single Chinese typeface.
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = clean_inline(part[2:-2] if bold else part)
        run = paragraph.add_run(value)
        set_run_font(run, size=size, bold=bold)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, rows, number):
    caption = doc.add_paragraph()
    configure_paragraph(caption, justified=False, first_indent=False, keep_next=True)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(caption, f"表{number}  {TABLE_TITLES[number - 1]}")

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_width = 6.27
    widths = [total_width / len(rows[0])] * len(rows[0])
    if len(rows[0]) == 3:
        widths = [1.08, 2.85, 2.34]
    elif len(rows[0]) == 4:
        widths = [1.18, 2.45, 1.45, 1.19]
    elif len(rows[0]) == 5:
        widths = [1.20, 1.17, 1.17, 1.17, 1.56]
    elif len(rows[0]) == 6:
        widths = [1.15, 1.02, 1.02, 0.98, 1.02, 1.08]
    elif len(rows[0]) == 7:
        widths = [1.12, 0.86, 0.86, 0.86, 0.78, 0.86, 0.93]

    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            add_rich_text(p, value)
            for run in p.runs:
                run.bold = r_idx == 0 or run.bold
    return table


def add_figure(doc, alt, relative_path, number):
    image_path = BASE_DIR / relative_path
    p = doc.add_paragraph()
    configure_paragraph(p, justified=False, first_indent=False, keep_next=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.85))

    caption = doc.add_paragraph()
    configure_paragraph(caption, justified=False, first_indent=False, keep_next=True)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(caption, f"图{number}  {alt}")

    interpretation = doc.add_paragraph()
    configure_paragraph(interpretation)
    add_rich_text(interpretation, "图形解读：" + FIGURE_INTERPRETATIONS[alt])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Title": (22, True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
        "Heading 1": (16, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
        "Heading 2": (14, True, WD_ALIGN_PARAGRAPH.LEFT, 9, 3),
        "Heading 3": (12, True, WD_ALIGN_PARAGRAPH.LEFT, 6, 0),
    }
    for name, (size, bold, align, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLUE if name != "Title" else RGBColor(0, 0, 0)
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(BODY_SIZE)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("机器学习量化交易策略研究"), size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    # Editorial-cover pattern, simplified for a formal course submission.
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    add_rich_text(title, "基于机器学习模型的量化交易策略", size=22)
    subtitle = doc.add_paragraph()
    configure_paragraph(subtitle, justified=False, first_indent=False)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(subtitle, "TASK6 课程作业", size=16)
    for _ in range(6):
        doc.add_paragraph()
    for label, value in (("姓名", "薛智鸣"), ("研究方法", "岭回归、决策树与随机森林"), ("完成日期", "2026年7月")):
        p = doc.add_paragraph()
        configure_paragraph(p, justified=False, first_indent=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_text(p, f"{label}：{value}", size=12)
    doc.add_page_break()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_number = 0
    figure_number = 0
    i = 1  # Skip source H1; cover replaces it.
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-:| ]+\|$", lines[i + 1].strip()):
            rows = []
            rows.append([clean_inline(x.strip()) for x in stripped.strip("|").split("|")])
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([clean_inline(x.strip()) for x in lines[i].strip().strip("|").split("|")])
                i += 1
            table_number += 1
            add_table(doc, rows, table_number)
            continue
        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            figure_number += 1
            add_figure(doc, image_match.group(1), image_match.group(2), figure_number)
            i += 1
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_rich_text(p, clean_inline(heading_match.group(2)), size={1: 16, 2: 14, 3: 12}[level])
            i += 1
            continue
        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            p = doc.add_paragraph(style="List Bullet" if bullet_match else "List Number")
            configure_paragraph(p, first_indent=False)
            add_rich_text(p, (bullet_match or number_match).group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        configure_paragraph(p)
        add_rich_text(p, stripped)
        i += 1

    # Keep Word metadata clean and explicit.
    doc.core_properties.title = "基于机器学习模型的量化交易策略"
    doc.core_properties.author = "薛智鸣"
    doc.core_properties.subject = "TASK6"
    doc.core_properties.keywords = "量化交易, 机器学习, 决策树, 随机森林, 回测"
    doc.save(DOCX_OUT)
    return DOCX_OUT


if __name__ == "__main__":
    print(build_document())
